import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib import colors

def generate_docx(md_path, docx_path):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []

    for line in lines:
        line_str = line.strip()
        
        if line_str.startswith('|') and line_str.endswith('|'):
            if '---' in line_str:
                continue
            cells = [c.strip() for c in line_str.split('|')[1:-1]]
            table_data.append(cells)
            in_table = True
            continue
        else:
            if in_table and table_data:
                cols = len(table_data[0])
                t = doc.add_table(rows=len(table_data), cols=cols)
                t.alignment = WD_TABLE_ALIGNMENT.CENTER
                t.style = 'Table Grid'
                for r_idx, row in enumerate(table_data):
                    for c_idx, val in enumerate(row):
                        cell = t.cell(r_idx, c_idx)
                        p = cell.paragraphs[0]
                        run = p.add_run(re.sub(r'\*\*(.*?)\*\*', r'\1', val))
                        if r_idx == 0:
                            run.bold = True
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
                table_data = []
                in_table = False

        if not line_str or line_str.startswith('<div') or line_str.startswith('</div>') or line_str == '<br>':
            continue

        if line_str.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(line_str[2:].replace('#', '').strip())
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(15, 23, 42)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line_str.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line_str[3:].strip())
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 41, 59)
        elif line_str.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line_str[4:].strip())
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(51, 65, 85)
        elif line_str.startswith('- ') or line_str.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = line_str[2:].strip()
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    p.add_run(part)
        elif re.match(r'^\d+\.\s', line_str):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s', '', line_str).strip()
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    p.add_run(part)
        elif line_str.startswith('```') or line_str.startswith('---'):
            continue
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line_str)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    p.add_run(part)

    if in_table and table_data:
        cols = len(table_data[0])
        t = doc.add_table(rows=len(table_data), cols=cols)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style = 'Table Grid'
        for r_idx, row in enumerate(table_data):
            for c_idx, val in enumerate(row):
                cell = t.cell(r_idx, c_idx)
                p = cell.paragraphs[0]
                run = p.add_run(re.sub(r'\*\*(.*?)\*\*', r'\1', val))
                if r_idx == 0:
                    run.bold = True

    doc.save(docx_path)
    print(f"DOCX Generated: {docx_path}")

def generate_pdf(md_path, pdf_path):
    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=A4,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CoverTitle',
        parent=styles['Heading1'],
        fontSize=18,
        leading=22,
        alignment=1,
        textColor=colors.HexColor('#0F172A'),
        spaceAfter=12
    )

    h1_style = ParagraphStyle(
        'H1Style',
        parent=styles['Heading1'],
        fontSize=13,
        leading=17,
        textColor=colors.HexColor('#1E293B'),
        spaceBefore=14,
        spaceAfter=8
    )

    h2_style = ParagraphStyle(
        'H2Style',
        parent=styles['Heading2'],
        fontSize=11,
        leading=15,
        textColor=colors.HexColor('#334155'),
        spaceBefore=10,
        spaceAfter=6
    )

    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        fontSize=9.5,
        leading=13.5,
        textColor=colors.HexColor('#1E293B'),
        spaceAfter=6
    )

    bullet_style = ParagraphStyle(
        'BulletStyle',
        parent=body_style,
        leftIndent=15,
        spaceAfter=4
    )

    story = []

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []

    for line in lines:
        line_str = line.strip()

        if line_str.startswith('|') and line_str.endswith('|'):
            if '---' in line_str:
                continue
            cells = [c.strip() for c in line_str.split('|')[1:-1]]
            table_data.append(cells)
            in_table = True
            continue
        else:
            if in_table and table_data:
                formatted_table = []
                for row_idx, row in enumerate(table_data):
                    formatted_row = []
                    for col in row:
                        clean_col = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', col)
                        clean_col = clean_col.replace('<br>', '<br/>')
                        if row_idx == 0:
                            clean_col = f"<b>{clean_col}</b>"
                        formatted_row.append(Paragraph(clean_col, body_style))
                    formatted_table.append(formatted_row)
                
                t = Table(formatted_table)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#F1F5F9')),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor('#0F172A')),
                    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 5),
                    ('TOPPADDING', (0,0), (-1,-1), 5),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
                ]))
                story.append(t)
                story.append(Spacer(1, 8))
                table_data = []
                in_table = False

        if not line_str or line_str.startswith('<div') or line_str.startswith('</div>') or line_str == '<br>':
            continue

        clean_line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line_str)
        clean_line = clean_line.replace('<br>', '<br/>')

        if line_str.startswith('# '):
            story.append(Paragraph(clean_line[2:].replace('#','').strip(), title_style))
            story.append(Spacer(1, 10))
        elif line_str.startswith('## '):
            story.append(Paragraph(clean_line[3:].strip(), h1_style))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#CBD5E1'), spaceAfter=8))
        elif line_str.startswith('### '):
            story.append(Paragraph(clean_line[4:].strip(), h2_style))
        elif line_str.startswith('- ') or line_str.startswith('* '):
            story.append(Paragraph("• " + clean_line[2:].strip(), bullet_style))
        elif re.match(r'^\d+\.\s', line_str):
            story.append(Paragraph(clean_line.strip(), bullet_style))
        elif line_str.startswith('```') or line_str.startswith('---'):
            continue
        else:
            story.append(Paragraph(clean_line, body_style))

    if in_table and table_data:
        formatted_table = []
        for row_idx, row in enumerate(table_data):
            formatted_row = []
            for col in row:
                clean_col = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', col)
                clean_col = clean_col.replace('<br>', '<br/>')
                formatted_row.append(Paragraph(clean_col, body_style))
            formatted_table.append(formatted_row)
        t = Table(formatted_table)
        t.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ]))
        story.append(t)

    doc.build(story)
    print(f"PDF Generated: {pdf_path}")

if __name__ == '__main__':
    base_dir = r"C:\Users\ArtvAdmirer\.gemini\antigravity\scratch\siresto"
    md = os.path.join(base_dir, "DOKUMEN_PERANCANGAN_PERANGKAT_LUNAK.md")
    docx = os.path.join(base_dir, "DOKUMEN_PERANCANGAN_PERANGKAT_LUNAK.docx")
    pdf = os.path.join(base_dir, "DOKUMEN_PERANCANGAN_PERANGKAT_LUNAK.pdf")
    
    generate_docx(md, docx)
    generate_pdf(md, pdf)
